from docx import Document
from io import BytesIO
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL


def format_response(response, header, division):
  doc = Document()

  # Header section
  if header:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.first_page_header

    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    run = paragraph.add_run()
    run.add_picture(f'static/{division}_header.png', width=section.page_width - section.left_margin - section.right_margin)

    doc.add_paragraph()
    
  # Title section
  title = doc.add_paragraph()
  title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
  title_run = title.add_run(response['resume']['candidate_name'].title())
  title_run.font.size = Pt(14)
  title_run.font.name = 'Segoe UI Bold'

  # Summary section
  summary = doc.add_paragraph()
  summary_run = summary.add_run("Summary")
  summary_run.font.size = Pt(12)
  summary_run.font.all_caps = True
  summary_run.font.name = 'Segoe UI Bold'

  summary_paragraph = doc.add_paragraph()
  summary_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
  summary_paragraph.paragraph_format.line_spacing = 1.5
  summary_run = summary_paragraph.add_run(response['resume']['summary'])
  summary_run.font.name = 'Segoe UI'

  # Skills section
  skills = doc.add_paragraph()
  skills_run = skills.add_run("Skills")
  skills_run.font.size = Pt(12)
  skills_run.font.all_caps = True
  skills_run.font.name = 'Segoe UI Bold'

  if len(response['resume']['skills']) == 0:
    skills_paragraph = doc.add_paragraph()
    skills_run = skills_paragraph.add_run("No skills listed.")
    skills_run.font.name = 'Segoe UI'
    skills_run.font.color.rgb = RGBColor(127, 127, 127)
    skills_run.font.italic = True
  else:
    for skill in response['resume']['skills']:
      skills_paragraph = doc.add_paragraph(style='List Bullet')
      skills_run = skills_paragraph.add_run(skill.title())
      skills_run.font.name = 'Segoe UI'

  doc.add_paragraph()
    
  # Education section
  education = doc.add_paragraph()
  education_run = education.add_run("Education")
  education_run.font.size = Pt(12)
  education_run.font.all_caps = True
  education_run.font.name = 'Segoe UI Bold'

  if len(response['resume']['education']) == 0:
    edu_paragraph = doc.add_paragraph()
    edu_run = edu_paragraph.add_run("No education listed.")
    edu_run.font.name = 'Segoe UI'
    edu_run.font.color.rgb = RGBColor(127, 127, 127)
    edu_run.font.italic = True
  else:
    for edu in response['resume']['education']:
      edu_school = doc.add_paragraph()
      edu_run = edu_school.add_run(edu['school'])
      edu_run.font.name = 'Segoe UI'
      edu_run.font.color.rgb = RGBColor(127, 127, 127)

      edu_degree = doc.add_paragraph()
      edu_run = edu_degree.add_run(f"{edu['degree']}, ({edu['date']})")
      edu_run.font.name = 'Segoe UI'

  doc.add_paragraph()

  # Experience section
  experience = doc.add_paragraph()
  experience_run = experience.add_run("Professional Experience")
  experience_run.font.size = Pt(12)
  experience_run.font.all_caps = True
  experience_run.font.name = 'Segoe UI Bold'

  for exp in response['resume']['experience']:
    work_dates = doc.add_paragraph()
    work_dates.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    work_dates_run = work_dates.add_run(f"{exp['employer']}\t\t\t\t{exp['work_dates']}")
    work_dates_run.font.name = 'Segoe UI'
    work_dates_run.font.color.rgb = RGBColor(127, 127, 127)

    job_title = doc.add_paragraph()
    job_title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    job_title_run = job_title.add_run(exp['job_title'])
    job_title_run.font.name = 'Segoe UI Bold'

    for r in exp['responsibilities']:
      job_responsibilities = doc.add_paragraph(style='List Bullet')
      job_responsibilities.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      job_responsibilities_run = job_responsibilities.add_run(r)
      job_responsibilities_run.font.name = 'Segoe UI'

    doc.add_paragraph()

  # Save the document to a BytesIO object
  doc_bin = BytesIO()
  doc.save(doc_bin)
  doc_bin.seek(0)
  
  return doc_bin